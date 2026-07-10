"""Word Document Parser using python-docx."""

from __future__ import annotations

from pathlib import Path
from docx import Document

from src.observability.logging import get_logger

logger = get_logger(__name__)


class DocxParser:
    """Parses text paragraphs and tabular structures in Microsoft Word documents."""

    @staticmethod
    def parse(file_path: Path | str) -> str:
        """Extracts text content from .docx file."""
        logger.info("Parsing DOCX file", path=str(file_path))
        try:
            doc = Document(file_path)
            extracted_text = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_text.append(" | ".join(row_text))

            return "\n\n".join(extracted_text).strip()
        except Exception as e:
            logger.error("DOCX parsing failed", error=str(e), path=str(file_path))
            return ""
